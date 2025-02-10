from docx import Document
from docx.shared import Pt  
from datetime import datetime
from docx2pdf import convert
from lxml import etree
import os
from django.conf import settings
from django.http import FileResponse


def filtrar_produtos(produtos, filtro):
    if filtro:
        if "-" in filtro:
            categoria, grupo = filtro.split("-")
            produtos = produtos.filter(grupo__slug=grupo, categoria__slug=categoria)
        else:
            produtos = produtos.filter(categoria__slug=filtro)
    return produtos

def ordenar_produtos(produtos, ordem):
    if ordem == "mais-orcados":
        lista_produtos = []
        for produto in produtos:
            lista_produtos.append((produto.total_orcados(), produto))
        lista_produtos = sorted(lista_produtos, reverse=True, key=lambda tupla: tupla[0])
        produtos = [item[1] for item in lista_produtos]
            
    return produtos

def exporta_dados_orcamento(endereco, orcamento):

    id_orcamento = orcamento.id
    total_orcamento = orcamento.quantidade_total
    codigo_website = orcamento.codigo_orcamento

    id_protheus = orcamento.cliente.protheus
    id_cliente = orcamento.cliente.id
    nome = orcamento.cliente.nome
    razaosocial = orcamento.cliente.razaosocial
    cnpjcpf = orcamento.cliente.cnpj
    email = orcamento.cliente.email
    telefone = orcamento.cliente.telefone

    id_endereco = endereco.id
    ruaavenida = endereco.rua
    numero_endereco = endereco.numero
    complemento = endereco.complemento
    cidade = endereco.cidade
    estado = endereco.estado
    cep = endereco.cep

    dicionario_valores = {
    "OOOO": str(id_orcamento).zfill(6),
    "WSWS": str(codigo_website),
    "DDD": str(datetime.now().day),
    "MMM": str(datetime.now().month),
    "AAA": str(datetime.now().year),
    "HMSS": str(datetime.now().strftime('%H:%M:%S')),
    "CCCC": str(id_protheus).zfill(6),
    "IIII": '0001',
    "RSRS": str(razaosocial),
    "EEEE": str(ruaavenida),
    "NNNN": str(numero_endereco),
    "CMCM": str(complemento),
    "CICI": str(cidade),
    "UFUF": str(estado),
    "PJPF": str(cnpjcpf),
    "CPCP": str(cep),
    "TFTF": str(telefone),
    "ELEL": '-',
    "OBOB": '',
    "TTTT": str(total_orcamento),
    "YYYY": str(datetime.now().year),
    }

    caminho_modelo = os.path.join(settings.MEDIA_ROOT, 'formularios', 'ModeloOrcamento-Metalforte.docx')
    caminho_saida = os.path.join(settings.MEDIA_ROOT, 'formularios', f'Orcamento-Metalforte-Nr-{id_orcamento}-{razaosocial}.docx')
    caminho_pdf = os.path.join(settings.MEDIA_ROOT, 'formularios', f'Orcamento-Metalforte-Nr-{id_orcamento}-{razaosocial}.pdf')
    
    orca = Document(caminho_modelo)

    for elemento in orca.element.body.iter():
        if etree.QName(elemento).localname == "t": 
            texto = elemento.text
            if texto:
                for codigo, valor in dicionario_valores.items():
                    if codigo in texto:
                        texto = texto.replace(codigo, valor)
                elemento.text = texto


    tabela = orca.tables[1] 


    for row in tabela.rows:
        for cell in row.cells:
            cell.text = ''


    for index, item in enumerate(orcamento.itens):
        sku = item.item_estoque.produto.SKU
        descricao_produto = item.item_estoque.produto.nome
        quantidade_solicitado = item.quantidade
        medida_produto = item.item_estoque.produto.medida.nome
   
        row_cells = tabela.add_row().cells
       
        row_cells[0].text = str(index + 1).zfill(2)  
        row_cells[1].text = str(sku).zfill(8)         
        row_cells[2].text = str(descricao_produto)    
        row_cells[3].text = ''                         
        row_cells[4].text = str(quantidade_solicitado)  
        row_cells[5].text = str(medida_produto)       
        
        for i in range(6, len(row_cells)):
            row_cells[i].text = '-'  

        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Arial MT'  
                    run.font.size = Pt(6.5)
    
    if index < 22:
        tabela = orca.tables[2]  


        for row in tabela.rows:
            for cell in row.cells:
                cell.text = ''

    orca.save(caminho_saida)

    convert(caminho_saida, caminho_pdf)

    print(caminho_pdf)
    
    return caminho_pdf

def exporta_formpdf(forms):
    
    for form in forms.values_list():
       caminho = form[1]

    formulario_nome = caminho.split("formularios\\")[-1] 

    resposta = FileResponse(open(caminho, 'rb'), as_attachment=True, filename=formulario_nome)
    return resposta